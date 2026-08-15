from __future__ import annotations

import re
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
INK = "202124"
MUTED = "667085"
GRID = "B8C4D1"
HEADER_FILL = "E8EEF5"
ALT_FILL = "F7F9FC"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"
TOTAL_DXA = 9360
TABLE_INDENT_DXA = 120
CJK_FONT = "等线"
LATIN_FONT = "Calibri"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color=GRID, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TOTAL_DXA:
        widths_dxa[-1] += TOTAL_DXA - sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TOTAL_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    new_grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        new_grid.append(col)
    tbl.replace(old_grid, new_grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def display_len(text: str) -> int:
    total = 0
    for ch in text:
        total += 2 if ord(ch) > 127 else 1
    return total


def calculate_widths(rows: list[list[str]]) -> list[int]:
    count = len(rows[0])
    if count == 1:
        return [TOTAL_DXA]
    if count == 2:
        first = max(display_len(r[0]) for r in rows)
        second = max(display_len(r[1]) for r in rows)
        ratio = max(0.20, min(0.38, first / max(first + second, 1)))
        return [round(TOTAL_DXA * ratio), TOTAL_DXA - round(TOTAL_DXA * ratio)]
    if count == 3:
        scores = []
        for i in range(count):
            scores.append(max(8, min(34, max(display_len(r[i]) for r in rows))))
        mins = [0.16, 0.18, 0.28]
    elif count == 4:
        scores = [max(7, min(28, max(display_len(r[i]) for r in rows))) for i in range(count)]
        mins = [0.12, 0.14, 0.20, 0.14]
    elif count == 5:
        scores = [max(6, min(24, max(display_len(r[i]) for r in rows))) for i in range(count)]
        mins = [0.10, 0.12, 0.18, 0.12, 0.12]
    else:
        scores = [max(5, min(18, max(display_len(r[i]) for r in rows))) for i in range(count)]
        mins = [0.08] * count

    total_min = sum(mins)
    remaining = 1.0 - total_min
    score_sum = sum(scores)
    fractions = [mins[i] + remaining * scores[i] / score_sum for i in range(count)]
    widths = [round(TOTAL_DXA * f) for f in fractions]
    widths[-1] += TOTAL_DXA - sum(widths)
    return widths


def set_run_font(run, latin=LATIN_FONT, east_asia=CJK_FONT, size=None, color=None,
                 bold=None, italic=None) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, latin=LATIN_FONT, east_asia=CJK_FONT, size=11, color=INK,
                   bold=False, italic=False) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    style.font.italic = italic
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        set_style_font(style, size=11, color=INK)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    caption = doc.styles["Caption"]
    set_style_font(caption, size=9, color=MUTED, italic=False)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def add_numbering_definition(doc: Document, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if not ordered:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_field(run, instruction: str, display: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True


def configure_header_footer(section, version: str, short_title: str = "完整功能规划") -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"连接器型商业收银经营平台    |    {short_title} {version}")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    r1 = fp.add_run("第 ")
    set_run_font(r1, size=8.5, color=MUTED)
    r2 = fp.add_run()
    set_run_font(r2, size=8.5, color=MUTED)
    add_field(r2, "PAGE", "1")
    r3 = fp.add_run(" 页")
    set_run_font(r3, size=8.5, color=MUTED)


def paragraph_bottom_border(paragraph, color=BLUE, size="12", space="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_callout_paragraph(paragraph, fill=HEADER_FILL, border="D5DEE8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), border)
        p_bdr.append(node)
    p_pr.append(p_bdr)


def add_cover(doc: Document, subtitle: str, status: str, technical: bool = False,
              specification: bool = False, document_no: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    if specification:
        kicker = "DETAILED DESIGN / DATA SPECIFICATION"
    elif technical:
        kicker = "TECHNICAL ARCHITECTURE / ENGINEERING STANDARD"
    else:
        kicker = "PRODUCT REQUIREMENTS / SYSTEM BLUEPRINT"
    r = p.add_run(kicker)
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("连接器型商业收银经营平台")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(subtitle)
    set_run_font(r, size=16, color=DARK_BLUE, bold=True)

    if specification:
        metadata = [
            ("文档编号", document_no or "待分配"),
            ("文档性质", "领域详细设计、数据库基线与研发验收依据"),
            ("服务端基线", "RuoYi-Vue-Plus 模块化单体 + MySQL 8.4 LTS"),
            ("终端基线", "Flutter Android POS，Windows认证后兼容"),
            ("文档状态", status),
            ("基线日期", "2026年8月15日"),
        ]
    elif technical:
        metadata = [
            ("文档性质", "技术架构基线、工程开发规范与质量门禁"),
            ("首发终端", "Flutter Android 商用POS一体机"),
            ("服务端", "RuoYi-Vue-Plus 模块化单体"),
            ("管理后台", "Vue 3 + TypeScript"),
            ("文档状态", status),
            ("基线日期", "2026年8月15日"),
        ]
    else:
        metadata = [
            ("文档性质", "产品立项基线、总体需求规格、架构边界与商用验收依据"),
            ("适用范围", "中国大陆零售门店、成长型连锁及相关本地生活业务"),
            ("核心模式", "核心经营平台 + 标准连接器平台 + 外部业务系统"),
            ("首个连接器", "鲸熵汇"),
            ("文档状态", status),
            ("更新日期", "2026年8月15日"),
        ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, size=10.5, color=INK, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(18)
    paragraph_bottom_border(p, color=BLUE, size="10", space="6")

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.left_indent = Inches(0.12)
    cp.paragraph_format.right_indent = Inches(0.12)
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after = Pt(4)
    style_callout_paragraph(cp)
    if specification:
        callout = "  数据边界先于代码，交易事实不可覆盖，所有关键结果必须可追溯和可核验。  "
    elif technical:
        callout = "  Android优先、Flutter统一终端、领域规则独立、可靠性门禁先行。  "
    else:
        callout = "  稳定收银为核心，连接器解耦外部生态，真实门店验收决定商业发布。  "
    cr = cp.add_run(callout)
    set_run_font(cr, size=12, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("jshPOS 项目组")
    set_run_font(r, size=10, color=MUTED, bold=True)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("目录")
    set_run_font(r, size=20, color=NAVY, bold=True)
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.space_after = Pt(0)
    run = toc_p.add_run()
    set_run_font(run, size=10.5, color=INK)
    add_field(run, 'TOC \\o "1-3" \\h \\z \\u', "目录将在打开或导出文档时自动更新")
    doc.add_page_break()


def generate_architecture_diagram(output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1500, 760
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    font_bold_path = Path(r"C:\Windows\Fonts\msyhbd.ttc")
    if not font_path.exists():
        font_path = Path(r"C:\Windows\Fonts\simhei.ttf")
    if not font_bold_path.exists():
        font_bold_path = font_path
    title_font = ImageFont.truetype(str(font_bold_path), 34)
    box_font = ImageFont.truetype(str(font_bold_path), 24)
    small_font = ImageFont.truetype(str(font_path), 19)

    draw.text((60, 35), "连接器型商业收银经营平台逻辑架构", font=title_font, fill="#17365D")

    def box(x1, y1, x2, y2, title, detail, fill):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=16, fill=fill, outline="#8AA4BF", width=3)
        tw = draw.textbbox((0, 0), title, font=box_font)[2]
        draw.text(((x1+x2-tw)/2, y1+18), title, font=box_font, fill="#17365D")
        lines = detail.split("\n")
        yy = y1 + 62
        for line in lines:
            lw = draw.textbbox((0, 0), line, font=small_font)[2]
            draw.text(((x1+x2-lw)/2, yy), line, font=small_font, fill="#344054")
            yy += 28

    def arrow(x1, y1, x2, y2):
        draw.line((x1, y1, x2, y2), fill="#667085", width=4)
        if x2 > x1:
            points = [(x2, y2), (x2-16, y2-9), (x2-16, y2+9)]
        else:
            points = [(x2, y2), (x2+16, y2-9), (x2+16, y2+9)]
        draw.polygon(points, fill="#667085")

    box(50, 145, 360, 590, "业务终端", "Flutter Android POS\nVue Web管理后台\nFlutter移动端 / PDA\nWindows POS（认证后）\n供应商与服务商端", "#F7F9FC")
    box(520, 115, 980, 620, "核心经营平台", "RuoYi-Vue-Plus平台底座\n商品 / 价格 / 促销\n订单 / 支付 / 售后\n库存 / 采购 / 供应链\n多租户 / 权限 / 审计", "#E8EEF5")
    box(1140, 145, 1450, 590, "连接器生态", "鲸熵汇\n美团 / 抖音\n京东 / 淘宝闪购\n配送 / 电子发票\n支付机构 / 金蝶用友", "#F7F9FC")
    arrow(360, 365, 520, 365)
    arrow(980, 365, 1140, 365)
    draw.text((510, 680), "连接器不得直写核心数据库；交易内核不依赖单一外部平台。", font=small_font, fill="#667085")
    img.save(output, quality=95)


def add_architecture_image(doc: Document, path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.35))
    doc_pr_nodes = run._r.xpath(".//wp:docPr")
    if doc_pr_nodes:
        doc_pr_nodes[0].set("descr", "连接器型商业收银经营平台逻辑架构图")
        doc_pr_nodes[0].set("title", "平台逻辑架构")
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("图1  连接器型商业收银经营平台逻辑架构")


def add_markdown_inline(paragraph, text: str, size=None, color=None) -> None:
    # Supports lightweight **bold** and `code` fragments.
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if size is not None or color is not None:
                set_run_font(run, size=size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, latin="Consolas", east_asia=CJK_FONT, size=(size or 10), color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            if size is not None or color is not None:
                set_run_font(run, size=size, color=color or INK)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    widths = calculate_widths(rows)
    set_table_geometry(table, widths)
    set_table_borders(table)
    font_size = 9.5 if cols <= 3 else 9 if cols <= 5 else 8
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        set_no_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=80 if cols <= 5 else 60,
                             bottom=80 if cols <= 5 else 60,
                             start=120,
                             end=120)
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, ALT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if cols >= 4 and (c_idx == 0 or value in {"P0", "P1", "P2", "P3", "待定"}):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_markdown_inline(p, value, size=font_size, color=INK)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = rgb(NAVY)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw.append(lines[i].strip())
        i += 1
    rows = []
    for idx, line in enumerate(raw):
        values = [x.strip() for x in line.strip("|").split("|")]
        if idx == 1 and all(re.fullmatch(r":?-{3,}:?", value or "") for value in values):
            continue
        rows.append(values)
    expected = len(rows[0])
    normalized = []
    for row in rows:
        if len(row) < expected:
            row += [""] * (expected - len(row))
        normalized.append(row[:expected])
    return normalized, i


def markdown_to_docx(md_path: Path, output_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    head = "\n".join(lines[:20])
    version_match = re.search(r"V\d+\.\d+", head)
    version = version_match.group(0) if version_match else "V1.0"
    subtitle = next((line[3:].strip() for line in lines[:20] if line.startswith("## ")), f"完整功能规划 {version}")
    status = next((line.split("：", 1)[1].strip().rstrip("  ") for line in lines[:20] if line.startswith("> 文档状态：")), "评审稿")
    document_no = next((line.split("：", 1)[1].strip().rstrip("  ") for line in lines[:20] if line.startswith("> 文档编号：")), "")
    technical = "技术架构与开发规范" in subtitle or "技术架构与开发规范" in head
    specification = "详细设计" in str(md_path.parent)
    short_title = subtitle.rsplit(" ", 1)[0] if specification else ("技术架构与开发规范" if technical else "完整功能规划")
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    configure_header_footer(doc.sections[0], version, short_title)
    bullet_num_id = add_numbering_definition(doc, ordered=False)
    decimal_num_id = add_numbering_definition(doc, ordered=True)
    add_cover(doc, subtitle, status, technical, specification, document_no)
    add_toc(doc)

    # Skip Markdown cover block through the first thematic divider.
    start = 0
    for idx, line in enumerate(lines):
        if line.strip() == "---":
            start = idx + 1
            break

    architecture_path = output_path.parent / ".docx_build" / "architecture.png"
    generate_architecture_diagram(architecture_path)

    i = start
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    architecture_inserted = False
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                if code_lang == "mermaid" and not architecture_inserted:
                    add_architecture_image(doc, architecture_path)
                    architecture_inserted = True
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.08)
                    p.paragraph_format.right_indent = Inches(0.08)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.05
                    style_callout_paragraph(p, fill=CALLOUT_FILL, border="D0D5DD")
                    r = p.add_run("\n".join(code_lines))
                    set_run_font(r, latin="Consolas", east_asia=CJK_FONT, size=8.5, color=INK)
                in_code = False
                code_lang = ""
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            p = doc.add_paragraph(style=f"Heading {level}")
            add_markdown_inline(p, text)
            i += 1
            continue

        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:].rstrip("  "))
                i += 1
            table = doc.add_table(rows=1, cols=1)
            set_table_geometry(table, [TOTAL_DXA])
            set_table_borders(table, color="D5DEE8")
            cell = table.cell(0, 0)
            set_cell_shading(cell, CALLOUT_FILL)
            set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_markdown_inline(p, " ".join(quote_lines), size=10.5, color=DARK_BLUE)
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            apply_numbering(p, bullet_num_id)
            add_markdown_inline(p, stripped[2:].strip())
            i += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, decimal_num_id)
            add_markdown_inline(p, number_match.group(2).strip())
            i += 1
            continue

        # Merge consecutive plain lines into one paragraph.
        parts = [stripped.rstrip("  ")]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if (nxt.startswith("#") or nxt.startswith("|") or nxt.startswith("- ") or
                    nxt.startswith("> ") or nxt.startswith("```") or nxt == "---" or
                    re.match(r"^\d+\.\s+", nxt)):
                break
            parts.append(nxt.rstrip("  "))
            i += 1
        p = doc.add_paragraph()
        add_markdown_inline(p, " ".join(parts))

    # Keep major headings on new pages but avoid a blank page after TOC caused by the first heading.
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    props = doc.core_properties
    props.title = f"连接器型商业收银经营平台：{subtitle}"
    props.subject = ("领域模型、数据库设计、数据边界与研发验收" if specification
                     else ("技术架构、开发规范、质量门禁与终端策略" if technical
                           else "产品立项、总体需求、架构边界与商用验收"))
    props.author = "jshPOS项目组"
    props.keywords = ("POS, 领域模型, 数据库, 账本, 多租户, ULID, Outbox"
                      if specification else
                      ("POS, Flutter, Android, RuoYi-Vue-Plus, 多租户, 离线同步, 开发规范"
                       if technical else
                       "POS, 收银系统, 连接器, 鲸熵汇, 零售SaaS, 供应链"))
    props.comments = (f"详细设计规格版本 {version}" if specification
                      else (f"技术架构与开发规范版本 {version}" if technical
                            else f"商业可落地复核与功能规划优化版本 {version}"))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_product_plan_docx.py input.md output.docx")
    markdown_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))


if __name__ == "__main__":
    main()

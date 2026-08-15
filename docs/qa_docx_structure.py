from pathlib import Path
import sys
from zipfile import ZipFile

from docx import Document


def main() -> None:
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).with_name("连接器型商业收银经营平台_完整功能规划_V2.0.docx")
    doc = Document(path)
    headings = [(p.style.name, p.text) for p in doc.paragraphs if p.style.name.startswith("Heading")]
    all_text = "\n".join(p.text for p in doc.paragraphs)
    if "31_领域模型与数据库设计说明书" in path.name:
        required = [
            "核心设计决策",
            "领域边界与数据所有权",
            "多租户、组织与门店模型",
            "商品、价格与促销数据模型",
            "订单、班次、支付与退款数据模型",
            "库存、成本、采购与调拨数据模型",
            "设备、同步与连接器数据模型",
            "核心表DDL基线",
            "数据库迁移与版本兼容",
            "数据质量与对账规则",
        ]
    elif "32_订单支付退款状态机规格" in path.name:
        required = [
            "总体设计原则",
            "销售订单状态机",
            "支付状态机",
            "退款与退货状态机",
            "命令、事件与数据契约",
            "事务边界与编排",
            "并发、幂等与数据库约束",
            "权限、审批与审计",
            "离线与弱网规则",
            "对账与异常处置",
            "验收测试",
        ]
    elif "33_库存账本预占与成本核算规格" in path.name:
        required = [
            "库存领域边界",
            "库存维度与数量口径",
            "库存流水",
            "预占模型",
            "负库存与超卖策略",
            "移动加权平均成本",
            "批次、效期与序列号",
            "盘点",
            "离线库存",
            "对账、重建与归档",
            "验收测试",
        ]
    elif "34_促销规则与优惠分摊规格" in path.name:
        required = [
            "核心原则",
            "金额与数量",
            "输入模型",
            "价格候选",
            "促销规则模型",
            "计算流水线",
            "优惠券、会员与预算",
            "优惠分摊",
            "退货与优惠回退",
            "促销快照与解释",
            "离线规则包",
            "验收测试",
        ]
    elif "35_POS离线同步协议" in path.name:
        required = [
            "设计原则",
            "组件与职责",
            "连接状态",
            "设备激活与会话",
            "本地数据库",
            "数据流与游标",
            "协议接口",
            "数据包协议",
            "冲突处理矩阵",
            "离线能力矩阵",
            "安全",
            "Schema 与版本兼容",
            "崩溃与灾难恢复",
            "验收测试",
        ]
    elif "36_Android设备适配协议与硬件认证手册" in path.name:
        required = [
            "适配范围",
            "总体架构",
            "能力模型",
            "统一调用协议",
            "打印协议",
            "扫码器协议",
            "电子秤协议",
            "Android 平台规范",
            "厂商 Adapter 规范",
            "硬件认证体系",
            "认证门禁与发布",
            "验收测试",
        ]
    elif "37_连接器SDK与标准数据契约" in path.name:
        required = [
            "连接器定位与边界",
            "运行架构",
            "连接器实例与生命周期",
            "Capability Manifest",
            "SDK SPI",
            "标准事件信封",
            "通用数据类型",
            "商品契约",
            "订单契约",
            "授权与密钥",
            "可靠性",
            "Schema 与版本",
            "连接器认证",
            "验收测试",
        ]
    elif "38_开放平台API与Webhook规范" in path.name:
        required = [
            "开放平台边界",
            "应用、安装与授权",
            "HTTP 与 URI",
            "标识、金额、数量与时间",
            "查询、分页与排序",
            "写入、幂等与并发",
            "响应与错误",
            "Webhook 订阅",
            "Webhook 签名",
            "API 版本与生命周期",
            "安全",
            "认证与发布",
            "验收测试",
        ]
    elif "39_安全隐私等保与灾备实施方案" in path.name:
        required = [
            "治理与责任",
            "等保实施路线",
            "资产与数据分类",
            "多租户安全",
            "身份与访问",
            "密码与密钥",
            "数据加密与脱敏",
            "支付安全边界",
            "应用与 API 安全",
            "隐私治理",
            "备份策略",
            "高可用与灾备",
            "事件响应",
            "安全验收",
        ]
    elif "40_商业V1验收测试计划" in path.name:
        required = [
            "验收范围",
            "验收组织",
            "准入与退出",
            "环境与数据",
            "需求追踪",
            "订单、支付与退款",
            "库存与成本",
            "促销",
            "POS 离线同步",
            "Android 与硬件",
            "连接器",
            "开放平台",
            "安全与隐私",
            "性能与容量",
            "备份与灾备",
            "实施与试营业",
            "Go/No-Go",
            "商业 V1 最终验收清单",
        ]
    elif "技术架构与开发规范" in path.name:
        required = [
            "技术选型结论",
            "总体架构",
            "服务端架构规范",
            "Flutter 客户端架构规范",
            "设备适配与 Edge Agent 规范",
            "多租户与身份安全规范",
            "数据库与领域数据规范",
            "API、事件与连接器规范",
            "编码规范",
            "测试与质量门禁",
            "AI 辅助开发规范",
            "必须建立的 ADR",
        ]
    else:
        required = [
            "数据主权与数据治理",
            "核心领域状态机与业务规则",
            "离线营业与同步规则",
            "非功能需求与质量指标",
            "设备与边缘服务中心",
            "开放平台",
            "商业发布总验收清单",
            "鲸熵汇连接器",
        ]
    if "V3.0" in path.name:
        required.extend(["标准实施包", "商业与单位经济指标", "管理层Go/No-Go决策表"])
    with ZipFile(path) as package:
        names = set(package.namelist())
        xml = package.read("word/document.xml").decode("utf-8")
        settings = package.read("word/settings.xml").decode("utf-8")
    report = {
        "docx_bytes": path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "headings": len(headings),
        "heading1": sum(1 for style, _ in headings if style == "Heading 1"),
        "heading2": sum(1 for style, _ in headings if style == "Heading 2"),
        "heading3": sum(1 for style, _ in headings if style == "Heading 3"),
        "required_sections": {item: any(item in text for _, text in headings) for item in required},
        "toc_field": "TOC \\o" in xml,
        "update_fields": "w:updateFields" in settings,
        "header": any(name.startswith("word/header") for name in names),
        "footer": any(name.startswith("word/footer") for name in names),
        "image_alt": "连接器型商业收银经营平台逻辑架构图" in xml,
        "replacement_char": "\ufffd" in all_text,
        "toc_placeholder_visible_in_model": "目录将在打开或导出文档时自动更新" in all_text,
    }
    print(report)
    print("FIRST_H1", [text for style, text in headings if style == "Heading 1"][:5])
    print("LAST_H1", [text for style, text in headings if style == "Heading 1"][-5:])
    assert all(report["required_sections"].values())
    assert report["toc_field"] and report["update_fields"]
    assert report["header"] and report["footer"]
    assert report["image_alt"] or report["inline_shapes"] == 0
    assert not report["replacement_char"]


if __name__ == "__main__":
    main()
